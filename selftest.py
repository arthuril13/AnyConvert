"""Round-trip test. Builds sample files, converts them, checks the results."""

import json
import shutil
import subprocess
import sys
import zipfile
from pathlib import Path

import engine
from engine import Tools, ConvertError

HERE = Path(__file__).resolve().parent
WORK = HERE / "_selftest"


def make_samples():
    if WORK.exists():
        shutil.rmtree(WORK)
    WORK.mkdir()
    s = {}

    from PIL import Image, ImageDraw
    im = Image.new("RGBA", (321, 187), (0, 0, 0, 0))       # odd size on purpose
    d = ImageDraw.Draw(im)
    d.ellipse([20, 20, 300, 160], fill=(255, 90, 60, 220))
    d.rectangle([120, 60, 260, 130], fill=(60, 140, 255, 255))
    s["png"] = WORK / "pic.png"
    im.save(s["png"])

    frames = []
    for i in range(6):
        f = Image.new("RGB", (120, 90), (20, 20, 30))
        ImageDraw.Draw(f).rectangle([i * 15, 20, i * 15 + 40, 70], fill=(240, 200, 60))
        frames.append(f)
    s["gif"] = WORK / "anim.gif"
    frames[0].save(s["gif"], save_all=True, append_images=frames[1:], duration=120, loop=0)

    s["txt"] = WORK / "notes.txt"
    s["txt"].write_text("Hello world.\nSecond line with an accent: café\n" * 20,
                        encoding="utf-8")

    s["md"] = WORK / "readme.md"
    s["md"].write_text("# Title\n\nSome **bold** text and a list:\n\n"
                       "- one\n- two\n- three\n\n## Section\n\n"
                       "```\ncode block\n```\n\nDone.\n", encoding="utf-8")

    s["csv"] = WORK / "table.csv"
    s["csv"].write_text('name,qty,price\nwidget,3,9.99\nbolt,120,0.35\n'
                        '"cog, with a comma",1,5.00\n', encoding="utf-8")

    s["json"] = WORK / "data.json"
    s["json"].write_text(json.dumps(
        [{"id": 1, "name": "alpha", "tags": ["x", "y"]},
         {"id": 2, "name": "beta", "tags": []}], indent=2), encoding="utf-8")

    from docx import Document
    doc = Document()
    doc.add_heading("Report", level=1)
    doc.add_paragraph("Intro paragraph with some words in it.")
    doc.add_heading("Findings", level=2)
    for t in ("first point", "second point"):
        doc.add_paragraph(t, style="List Bullet")
    t = doc.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "key"
    t.cell(0, 1).text = "value"
    t.cell(1, 0).text = "speed"
    t.cell(1, 1).text = "fast"
    s["docx"] = WORK / "report.docx"
    doc.save(str(s["docx"]))

    s["zip"] = WORK / "bundle.zip"
    with zipfile.ZipFile(s["zip"], "w") as z:
        z.writestr("root.txt", "top level file\n")
        z.writestr("sub/deep/inner file.txt", "nested content\n")
        z.write(str(s["png"]), "images/pic.png")

    if Tools.ffmpeg:
        s["wav"] = WORK / "tone.wav"
        subprocess.run([Tools.ffmpeg, "-y", "-loglevel", "error", "-f", "lavfi",
                        "-i", "sine=frequency=440:duration=2", str(s["wav"])],
                       check=True, creationflags=engine.NO_WINDOW)
        s["mp4"] = WORK / "clip.mp4"
        subprocess.run([Tools.ffmpeg, "-y", "-loglevel", "error", "-f", "lavfi",
                        "-i", "testsrc=size=320x240:rate=15:duration=2", "-f", "lavfi",
                        "-i", "sine=frequency=330:duration=2", "-c:v", "libx264",
                        "-pix_fmt", "yuv420p", "-c:a", "aac", "-strict", "-2",
                        "-shortest", str(s["mp4"])],
                       check=True, creationflags=engine.NO_WINDOW)

    fontdir = Path("C:/Windows/Fonts")
    for name in ("arial.ttf", "segoeui.ttf", "tahoma.ttf"):
        if (fontdir / name).is_file():
            s["ttf"] = WORK / "sample.ttf"
            shutil.copyfile(str(fontdir / name), str(s["ttf"]))
            break
    return s


CASES = [
    ("png", "jpg"), ("png", "webp"), ("png", "ico"), ("png", "bmp"), ("png", "tiff"),
    ("png", "pdf"), ("png", "gif"), ("png", "avif"),
    ("gif", "webp"), ("gif", "png"), ("gif", "mp4"),
    ("txt", "pdf"), ("txt", "html"), ("txt", "docx"),
    ("md", "html"), ("md", "pdf"), ("md", "docx"), ("md", "txt"),
    ("csv", "xlsx"), ("csv", "json"), ("csv", "html"), ("csv", "md"), ("csv", "tsv"),
    ("csv", "pdf"),
    ("json", "csv"), ("json", "yaml"), ("json", "xml"), ("json", "xlsx"),
    ("docx", "txt"), ("docx", "md"), ("docx", "pdf"), ("docx", "html"),
    ("zip", "7z"), ("zip", "tar.gz"), ("zip", "iso"),
    ("wav", "mp3"), ("wav", "flac"), ("wav", "ogg"), ("wav", "m4a"), ("wav", "opus"),
    ("mp4", "mp3"), ("mp4", "wav"), ("mp4", "avi"), ("mp4", "webm"), ("mp4", "mkv"),
    ("mp4", "gif"), ("mp4", "png"),
    ("ttf", "woff"), ("ttf", "woff2"),
]

# Second pass: convert an output again, to exercise the formats we just made.
CHAINED = [
    ("pic.pdf", "png"), ("pic.pdf", "txt", "expected to fail: image-only pdf"), ("readme.pdf", "txt"),
    ("table.xlsx", "csv"), ("table.xlsx", "json"), ("bundle.iso", "zip"),
    ("tone.mp3", "wav"), ("clip.avi", "mp3"),
]


def run():
    samples = make_samples()
    print("samples:", ", ".join(sorted(samples)))
    print()
    ok = fail = skip = 0
    fails = []

    def attempt(src, target, note=""):
        nonlocal ok, fail, skip
        if not Path(src).is_file():
            skip += 1
            print("  SKIP  %-28s -> %-6s (no sample)" % (Path(src).name, target))
            return None
        try:
            out = engine.convert(src, target, overwrite=True)
            size = out.stat().st_size
            if size == 0:
                raise ConvertError("output is empty")
            ok += 1
            print("  ok    %-28s -> %-6s %8d bytes  %s"
                  % (Path(src).name, target, size, note))
            return out
        except Exception as e:
            fail += 1
            fails.append("%s -> %s: %s" % (Path(src).name, target, e))
            print("  FAIL  %-28s -> %-6s %s" % (Path(src).name, target, e))
            return None

    print("direct conversions")
    for src_key, target in CASES:
        src = samples.get(src_key)
        if src is None:
            skip += 1
            print("  SKIP  .%-27s -> %-6s (no sample)" % (src_key, target))
            continue
        attempt(src, target)

    print()
    print("second pass over the results")
    for entry in CHAINED:
        name, target = entry[0], entry[1]
        if len(entry) > 2:
            print('  note  %-28s -> %-6s %s' % (name, target, entry[2]))
            continue
        attempt(WORK / name, target)

    print()
    print("content checks")
    checks = []

    p = WORK / "table.json"
    if p.is_file():
        rows = json.loads(p.read_text(encoding="utf-8"))
        checks.append(("csv -> json keeps the quoted comma",
                       rows[2]["name"] == "cog, with a comma"))
    p = WORK / "report.md"
    if p.is_file():
        t = p.read_text(encoding="utf-8")
        checks.append(("docx -> md keeps headings", "# Report" in t))
        checks.append(("docx -> md keeps the table", "speed" in t))
    p = WORK / "pic.ico"
    if p.is_file():
        from PIL import Image
        checks.append(("ico is a real icon", Image.open(p).size[0] in (16, 32, 48, 64, 128, 256)))
    p = WORK / "anim.webp"
    if p.is_file():
        from PIL import Image
        checks.append(("animation survives gif -> webp",
                       getattr(Image.open(p), "n_frames", 1) > 1))
    p = WORK / "bundle.iso"
    if p.is_file() and Tools.sevenzip:
        listing = engine.run([Tools.sevenzip, "l", "-ba", str(p)])
        checks.append(("iso contains the nested file", "inner" in listing.lower()))
    p = WORK / "clip.mp4.gif" if (WORK / "clip.mp4.gif").is_file() else WORK / "clip.gif"
    if p.is_file():
        from PIL import Image
        checks.append(("video -> gif is animated", getattr(Image.open(p), "n_frames", 1) > 1))
    p = WORK / "pic.jpg"
    if p.is_file():
        from PIL import Image
        checks.append(("transparent png -> jpg is flattened", Image.open(p).mode == "RGB"))

    for name, passed in checks:
        print("  %s  %s" % ("ok  " if passed else "FAIL", name))
        if passed:
            ok += 1
        else:
            fail += 1
            fails.append(name)

    print()
    print("=" * 62)
    print("passed %d   failed %d   skipped %d" % (ok, fail, skip))
    if fails:
        print()
        for f in fails:
            print("  !", f)
    return 1 if fail else 0


if __name__ == "__main__":
    sys.exit(run())
