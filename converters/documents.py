"""Documents: PDF, Office files, ebooks, Markdown, HTML and plain text."""

import os
import re
import html as htmlmod
import shutil
import tempfile
from html.parser import HTMLParser
from pathlib import Path

from engine import rule, run, Tools, ConvertError

TEXT_IN = "txt text log md markdown rst html htm xhtml json csv ini cfg conf srt vtt"


# --------------------------------------------------------------------------
# LibreOffice - by far the best route for Office files when it is installed
# --------------------------------------------------------------------------

LO_TEXT_IN = "doc docx odt rtf txt html htm fodt sxw uot wps abw"
LO_TEXT_OUT = "docx odt rtf txt html pdf fodt epub"
LO_SHEET_IN = "xls xlsx xlsm ods csv tsv fods dif slk dbf"
LO_SHEET_OUT = "xlsx xls ods csv html pdf fods"
LO_PRES_IN = "ppt pptx odp fodp sxi pps ppsx"
LO_PRES_OUT = "pptx odp pdf html"
LO_DRAW_IN = "odg vsd vsdx wmf emf cdr pub svg"
LO_DRAW_OUT = "pdf svg png odg"


def _soffice(src, out, dst):
    if not Tools.soffice:
        raise ConvertError("LibreOffice not found")
    workdir = Path(tempfile.mkdtemp(prefix="anyconv_lo_"))
    profile = workdir / "profile"
    try:
        run([Tools.soffice,
             "-env:UserInstallation=file:///" + str(profile).replace("\\", "/"),
             "--headless", "--norestore", "--convert-to", dst,
             "--outdir", str(workdir), str(src)], timeout=900)
        made = [p for p in workdir.iterdir()
                if p.is_file() and p.suffix.lower() == "." + dst]
        if not made:
            raise ConvertError("LibreOffice could not write a .%s from that file" % dst)
        shutil.move(str(made[0]), str(out))
    finally:
        shutil.rmtree(workdir, ignore_errors=True)


for _srcs, _dsts, _name in (
    (LO_TEXT_IN, LO_TEXT_OUT, "office text"),
    (LO_SHEET_IN, LO_SHEET_OUT, "office sheet"),
    (LO_PRES_IN, LO_PRES_OUT, "office slides"),
    (LO_DRAW_IN, LO_DRAW_OUT, "office drawing"),
):
    rule(_srcs, _dsts, need=("soffice",), cost=4, label=_name)(_soffice)


# --------------------------------------------------------------------------
# PDF and ebooks via PyMuPDF
# --------------------------------------------------------------------------

try:
    import pymupdf
except ImportError:                                   # older wheel name
    try:
        import fitz as pymupdf
    except ImportError:
        pymupdf = None

if pymupdf:
    EBOOK_IN = "epub mobi fb2 xps oxps cbz"

    @rule("pdf " + EBOOK_IN, "txt", cost=6, label="pdf text")
    def pdf_to_txt(src, out, dst):
        doc = pymupdf.open(str(src))
        parts = [page.get_text() for page in doc]
        pages = len(doc)
        doc.close()
        text = "\n\n".join(parts)
        if not text.strip():
            raise ConvertError(
                "there is no text in this file - its %d page(s) are scanned "
                "images. Convert it to PNG instead, or OCR it first." % pages)
        Path(out).write_text(text, encoding="utf-8")

    @rule("pdf " + EBOOK_IN, "html", cost=6, label="pdf html")
    def pdf_to_html(src, out, dst):
        doc = pymupdf.open(str(src))
        body = "\n".join(page.get_text("html") for page in doc)
        doc.close()
        Path(out).write_text(
            "<!doctype html>\n<meta charset='utf-8'>\n<title>%s</title>\n%s"
            % (htmlmod.escape(Path(src).stem), body), encoding="utf-8")

    @rule("pdf " + EBOOK_IN, "png jpg jpeg webp bmp tif tiff", cost=6, label="pdf pages")
    def pdf_to_images(src, out, dst):
        """Page 1 goes to the chosen name, further pages sit beside it."""
        doc = pymupdf.open(str(src))
        out = Path(out)
        mat = pymupdf.Matrix(150 / 72.0, 150 / 72.0)
        for i, page in enumerate(doc):
            pix = page.get_pixmap(matrix=mat)
            target = out if i == 0 else out.with_name("%s_p%d%s" % (out.stem, i + 1, out.suffix))
            if dst in ("jpg", "jpeg", "bmp") and pix.alpha:
                pix = pymupdf.Pixmap(pymupdf.csRGB, pix)
            pix.save(str(target))
        doc.close()

    @rule("pdf", "svg", cost=6, label="pdf to svg")
    def pdf_to_svg(src, out, dst):
        doc = pymupdf.open(str(src))
        Path(out).write_text(doc[0].get_svg_image(), encoding="utf-8")
        doc.close()

    @rule(EBOOK_IN, "pdf", cost=6, label="ebook to pdf")
    def ebook_to_pdf(src, out, dst):
        doc = pymupdf.open(str(src))
        pdf = doc.convert_to_pdf()
        Path(out).write_bytes(pdf)
        doc.close()


# --------------------------------------------------------------------------
# Calibre for real ebook output
# --------------------------------------------------------------------------

CAL_IN = "epub mobi azw azw3 azw4 fb2 lit lrf pdb pml rb snb tcr txt html htm rtf docx pdf cbz cbr"
CAL_OUT = "epub mobi azw3 fb2 lit lrf pdb pml rb snb tcr txt html rtf docx pdf"


@rule(CAL_IN, CAL_OUT, need=("calibre",), cost=4, label="ebook")
def calibre_convert(src, out, dst):
    run([Tools.calibre, str(src), str(out)], timeout=1800)


# --------------------------------------------------------------------------
# Word documents without LibreOffice
# --------------------------------------------------------------------------

try:
    import docx as pydocx
    from docx import Document
except ImportError:
    pydocx = None

if pydocx:
    def _docx_blocks(src):
        doc = Document(str(src))
        for p in doc.paragraphs:
            yield p.style.name if p.style is not None else "Normal", p.text
        for t in doc.tables:
            for row in t.rows:
                yield "Table", " | ".join(c.text.strip() for c in row.cells)

    @rule("docx", "txt", cost=8, label="docx text")
    def docx_to_txt(src, out, dst):
        lines = [text for _, text in _docx_blocks(src)]
        Path(out).write_text("\n".join(lines), encoding="utf-8")

    @rule("docx", "md", cost=8, label="docx markdown")
    def docx_to_md(src, out, dst):
        lines = []
        for style, text in _docx_blocks(src):
            if not text.strip():
                lines.append("")
            elif style.startswith("Heading"):
                level = "".join(c for c in style if c.isdigit()) or "1"
                lines.append("#" * min(int(level), 6) + " " + text)
            elif style.startswith("List"):
                lines.append("- " + text)
            elif style == "Table":
                lines.append("| " + text + " |")
            else:
                lines.append(text)
        Path(out).write_text("\n".join(lines) + "\n", encoding="utf-8")

    @rule("docx", "pdf", cost=7, label="docx to pdf")
    def docx_to_pdf(src, out, dst):
        """Without LibreOffice, go through Markdown so headings and lists survive."""
        if "text_to_pdf" not in globals():
            raise ConvertError("PDF output needs the reportlab package")
        tmp = Path(tempfile.mkdtemp(prefix="anyconv_dx_"))
        try:
            mid = tmp / "doc.md"
            docx_to_md(src, mid, "md")
            text_to_pdf(mid, out, "pdf")
        finally:
            shutil.rmtree(tmp, ignore_errors=True)


    @rule("txt md markdown", "docx", cost=8, label="text to docx")
    def txt_to_docx(src, out, dst):
        doc = Document()
        for raw in Path(src).read_text(encoding="utf-8", errors="replace").splitlines():
            line = raw.rstrip()
            m = re.match(r"^(#{1,6})\s+(.*)", line)
            if m:
                doc.add_heading(m.group(2), level=len(m.group(1)))
            elif re.match(r"^\s*[-*+]\s+", line):
                doc.add_paragraph(re.sub(r"^\s*[-*+]\s+", "", line), style="List Bullet")
            else:
                doc.add_paragraph(line)
        doc.save(str(out))


# --------------------------------------------------------------------------
# Markdown / HTML / text
# --------------------------------------------------------------------------

_CSS = ("body{font:16px/1.6 -apple-system,Segoe UI,Roboto,sans-serif;"
        "max-width:44em;margin:2.5em auto;padding:0 1.2em;color:#1a1a1a}"
        "pre,code{font-family:Consolas,monospace;background:#f4f4f5;border-radius:4px}"
        "pre{padding:.8em;overflow-x:auto}code{padding:.15em .35em}"
        "table{border-collapse:collapse}td,th{border:1px solid #ccc;padding:.4em .7em}"
        "blockquote{border-left:3px solid #ddd;margin-left:0;padding-left:1em;color:#555}")


def _page(title, body):
    return ("<!doctype html>\n<html><head><meta charset='utf-8'>\n"
            "<title>%s</title>\n<style>%s</style>\n</head><body>\n%s\n</body></html>\n"
            % (htmlmod.escape(title), _CSS, body))


try:
    import markdown as mdlib

    @rule("md markdown mdown mkd", "html", cost=6, label="markdown to html")
    def md_to_html(src, out, dst):
        text = Path(src).read_text(encoding="utf-8", errors="replace")
        body = mdlib.markdown(text, extensions=["tables", "fenced_code", "toc", "sane_lists"])
        Path(out).write_text(_page(Path(src).stem, body), encoding="utf-8")

except ImportError:
    pass


class _Stripper(HTMLParser):
    SKIP = {"script", "style", "head", "meta", "link"}
    BLOCK = {"p", "div", "br", "li", "tr", "h1", "h2", "h3", "h4", "h5", "h6",
             "section", "article", "blockquote", "pre", "table"}

    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.out, self._skip = [], 0

    def handle_starttag(self, tag, attrs):
        if tag in self.SKIP:
            self._skip += 1
        elif tag in self.BLOCK:
            self.out.append("\n")
        elif tag == "img":
            alt = dict(attrs).get("alt")
            if alt:
                self.out.append("[%s]" % alt)

    def handle_endtag(self, tag):
        if tag in self.SKIP and self._skip:
            self._skip -= 1
        elif tag in self.BLOCK:
            self.out.append("\n")

    def handle_data(self, data):
        if not self._skip:
            self.out.append(data)

    def text(self):
        t = "".join(self.out)
        t = re.sub(r"[ \t]+", " ", t)
        t = re.sub(r"\n\s*\n\s*\n+", "\n\n", t)
        return "\n".join(l.strip() for l in t.splitlines()).strip() + "\n"


@rule("html htm xhtml", "txt", cost=6, label="html to text")
def html_to_txt(src, out, dst):
    p = _Stripper()
    p.feed(Path(src).read_text(encoding="utf-8", errors="replace"))
    Path(out).write_text(p.text(), encoding="utf-8")


@rule("html htm xhtml", "md", cost=7, label="html to markdown")
def html_to_md(src, out, dst):
    raw = Path(src).read_text(encoding="utf-8", errors="replace")
    raw = re.sub(r"(?is)<(script|style|head)[^>]*>.*?</\1>", "", raw)
    for i in range(1, 7):
        raw = re.sub(r"(?is)<h%d[^>]*>(.*?)</h%d>" % (i, i),
                     lambda m, i=i: "\n\n%s %s\n\n" % ("#" * i, m.group(1).strip()), raw)
    raw = re.sub(r"(?is)<(b|strong)[^>]*>(.*?)</\1>", r"**\2**", raw)
    raw = re.sub(r"(?is)<(i|em)[^>]*>(.*?)</\1>", r"*\2*", raw)
    raw = re.sub(r"(?is)<code[^>]*>(.*?)</code>", r"`\1`", raw)
    raw = re.sub(r"(?is)<a[^>]*href=[\"']([^\"']+)[\"'][^>]*>(.*?)</a>", r"[\2](\1)", raw)
    raw = re.sub(r"(?is)<li[^>]*>(.*?)</li>", r"\n- \1", raw)
    raw = re.sub(r"(?is)<(p|div|tr|br\s*/?)[^>]*>", "\n\n", raw)
    p = _Stripper()
    p.feed(raw)
    Path(out).write_text(p.text(), encoding="utf-8")


@rule("txt text log rst srt vtt ini cfg conf", "html", cost=8, label="text to html")
def txt_to_html(src, out, dst):
    text = Path(src).read_text(encoding="utf-8", errors="replace")
    Path(out).write_text(_page(Path(src).stem, "<pre>%s</pre>" % htmlmod.escape(text)),
                         encoding="utf-8")


@rule("md markdown mdown mkd", "txt", cost=9, label="markdown to text")
def md_to_txt(src, out, dst):
    shutil.copyfile(str(src), str(out))


# --------------------------------------------------------------------------
# Anything text-ish to PDF, via ReportLab (no LibreOffice needed)
# --------------------------------------------------------------------------

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Preformatted, Spacer

    _FONT, _MONO = "Helvetica", "Courier"
    for _name, _file, _target in (("ACSans", "segoeui.ttf", "body"),
                                  ("ACMono", "consola.ttf", "mono")):
        _p = Path(os.environ.get("WINDIR", r"C:\Windows")) / "Fonts" / _file
        if _p.is_file():
            try:
                pdfmetrics.registerFont(TTFont(_name, str(_p)))
                if _target == "body":
                    _FONT = _name
                else:
                    _MONO = _name
            except Exception:
                pass

    @rule("txt text log rst srt vtt ini cfg conf md markdown html htm xhtml csv json",
          "pdf", cost=8, label="text to pdf")
    def text_to_pdf(src, out, dst):
        ext = Path(src).suffix.lower().lstrip(".")
        if ext in ("html", "htm", "xhtml"):
            p = _Stripper()
            p.feed(Path(src).read_text(encoding="utf-8", errors="replace"))
            text = p.text()
        else:
            text = Path(src).read_text(encoding="utf-8", errors="replace")

        styles = getSampleStyleSheet()
        body = ParagraphStyle("acbody", parent=styles["Normal"], fontName=_FONT,
                              fontSize=10.5, leading=15)
        mono = ParagraphStyle("acmono", parent=styles["Code"], fontName=_MONO,
                              fontSize=8.5, leading=11)
        heads = {n: ParagraphStyle("ach%d" % n, parent=body, fontName=_FONT,
                                   fontSize=20 - 2 * n, leading=24 - 2 * n,
                                   spaceBefore=10, spaceAfter=4)
                 for n in range(1, 7)}

        flow, fenced, buf = [], False, []
        for raw in text.splitlines():
            if raw.strip().startswith("```"):
                if fenced:
                    flow.append(Preformatted("\n".join(buf), mono))
                    buf = []
                fenced = not fenced
                continue
            if fenced:
                buf.append(raw)
                continue
            line = raw.rstrip()
            if not line:
                flow.append(Spacer(1, 5))
                continue
            m = re.match(r"^(#{1,6})\s+(.*)", line)
            if m:
                flow.append(Paragraph(htmlmod.escape(m.group(2)), heads[len(m.group(1))]))
            elif line.startswith(("    ", "\t")):
                flow.append(Preformatted(line, mono))
            else:
                flow.append(Paragraph(htmlmod.escape(line), body))
        if buf:
            flow.append(Preformatted("\n".join(buf), mono))

        SimpleDocTemplate(str(out), pagesize=A4, title=Path(src).stem,
                          leftMargin=20 * mm, rightMargin=20 * mm,
                          topMargin=18 * mm, bottomMargin=18 * mm).build(flow or [Spacer(1, 1)])

except ImportError:
    pass
